#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Unified toolkit for exporting chat logs, training causal LMs, and running utilities."""
from __future__ import annotations

import argparse
import json
import logging
import math
import os
import re
import textwrap
import time
from dataclasses import dataclass
from datetime import datetime
from string import ascii_letters, digits
from typing import Iterable, List, Optional

logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')

# ============================================================
# القسم (أ): مُصدِّر سجلات الدردشة
# ============================================================
SUPPORTED_ROLES = {"user": "المستخدم", "assistant": "المساعد", "system": "النظام"}


@dataclass
class Message:
    role: str
    content: str
    timestamp: Optional[str] = None  # ISO8601 إذا متوفر


def load_messages(path: str) -> List[Message]:
    import csv

    ext = os.path.splitext(path)[1].lower()
    msgs: List[Message] = []
    if ext == ".csv":
        with open(path, newline='', encoding="utf-8") as f:
            sample = f.read(4096)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample)
            except csv.Error:
                dialect = csv.excel
            reader = csv.DictReader(f, dialect=dialect)
            for r in reader:
                role = (r.get("role") or "").strip()
                content = (r.get("content") or "")
                ts = r.get("timestamp") or r.get("time")
                if role or content:
                    msgs.append(Message(role=role, content=content, timestamp=ts))
    elif ext == ".jsonl":
        with open(path, encoding="utf-8") as f:
            for i, line in enumerate(f, 1):
                if not line.strip():
                    continue
                try:
                    o = json.loads(line)
                except json.JSONDecodeError as e:
                    logging.warning("JSONL غير صالح في السطر %d: %s", i, e)
                    continue
                msgs.append(
                    Message(
                        role=(o.get("role") or ""),
                        content=(o.get("content") or ""),
                        timestamp=o.get("timestamp") or o.get("time"),
                    )
                )
    elif ext == ".json":
        with open(path, encoding="utf-8") as f:
            try:
                data = json.load(f)
            except json.JSONDecodeError as e:
                raise ValueError(f"JSON غير صالح: {e}")
        items = data.get("messages") if isinstance(data, dict) else data
        if not isinstance(items, list):
            raise ValueError("ملف JSON يجب أن يكون قائمة رسائل أو كائنًا يحتوي على مفتاح 'messages'.")
        for i, o in enumerate(items, 1):
            if not isinstance(o, dict):
                logging.warning("عنصر غير قاموسي عند %d، سيتم تجاوزه.", i)
                continue
            msgs.append(
                Message(
                    role=(o.get("role") or ""),
                    content=(o.get("content") or ""),
                    timestamp=o.get("timestamp") or o.get("time"),
                )
            )
    else:
        raise ValueError("تنسيق الإدخال غير مدعوم. استخدم .csv أو .jsonl أو .json")
    if not msgs:
        logging.warning("لا توجد رسائل صالحة للمعالجة.")
    return msgs


def _local_iso() -> str:
    return datetime.now().astimezone().isoformat(timespec="seconds")


MD_SPECIALS = ("*", "_", "#", ">", "`")


def md_escape_first_char(s: str) -> str:
    out_lines = []
    for ln in s.replace("\r", "").split("\n"):
        if ln and ln[0] in MD_SPECIALS:
            out_lines.append("\\" + ln)
        else:
            out_lines.append(ln)
    return "\n".join(out_lines)


def to_markdown(msgs: List[Message], title: str = "سجل الدردشة") -> str:
    lines = [f"# {title}", f"_تم التصدير: {_local_iso()}_\n"]
    for i, m in enumerate(msgs, 1):
        role_key = (m.role or "").lower()
        prefix = f"**{SUPPORTED_ROLES.get(role_key, m.role or 'غير معروف')}**"
        ts = f"\n> _{m.timestamp}_" if getattr(m, "timestamp", None) else ""
        lines.append(f"### {i}. {prefix}{ts}")
        lines.append(md_escape_first_char(m.content) + "\n")
    return "\n".join(lines)


def save_markdown(md: str, outpath: str) -> None:
    with open(outpath, "w", encoding="utf-8") as f:
        f.write(md)


def save_docx(msgs: List[Message], outpath: str, title: str = "سجل الدردشة", font_name: str = "Arial") -> None:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    def set_paragraph_rtl(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    doc = Document()
    doc.add_heading(title, 0)
    style = doc.styles["Normal"]
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    style.font.size = Pt(11)

    for i, m in enumerate(msgs, 1):
        role_key = (m.role or "").lower()
        role_title = SUPPORTED_ROLES.get(role_key, m.role or 'غير معروف')
        h = doc.add_heading(f"{i}. {role_title}", level=2)
        set_paragraph_rtl(h)
        p = doc.add_paragraph(m.content)
        set_paragraph_rtl(p)

    doc.save(outpath)


def save_xlsx(msgs: List[Message], outpath: str) -> None:
    import pandas as pd

    df = pd.DataFrame([{"role": m.role, "content": m.content, "timestamp": m.timestamp} for m in msgs])
    df.index = df.index + 1
    df.index.name = "idx"
    df.to_excel(outpath, index=True)


def save_pdf(md_text: str, outpath: str, font_paths: Optional[Iterable[str]] = None) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    width, height = A4
    c = canvas.Canvas(outpath, pagesize=A4)

    chosen_font = "Helvetica"

    def reshape_default(s: str) -> str:
        return s

    reshape = reshape_default
    try:
        import arabic_reshaper  # type: ignore
        from bidi.algorithm import get_display  # type: ignore

        def reshape(s: str) -> str:  # type: ignore[redefined-outer-name]
            try:
                return get_display(arabic_reshaper.reshape(s))
            except Exception:
                return s
    except Exception:
        logging.warning("لم يتم العثور على arabic-reshaper/python-bidi؛ سيُطبع النص دون تشكيل RTL.")

    candidates = [
        "Amiri-Regular.ttf",
        "Amiri.ttf",
        "Arial.ttf",
        "NotoNaskhArabic-Regular.ttf",
        "NotoSansArabic-Regular.ttf",
    ]
    if font_paths:
        candidates = list(font_paths) + candidates

    for fp in candidates:
        if os.path.isfile(fp):
            try:
                pdfmetrics.registerFont(TTFont("CustomArabic", fp))
                chosen_font = "CustomArabic"
                break
            except Exception:
                continue

    c.setFont(chosen_font, 12)

    x_left = 20 * mm
    x_right = width - 20 * mm
    y = height - 20 * mm
    line_h = 6 * mm
    wrap_chars = 95

    for raw_para in md_text.split("\n"):
        para = raw_para.rstrip("\n")
        lines = textwrap.wrap(para, width=wrap_chars) if para.strip() else [""]
        for line in lines:
            s = reshape(line)
            if s and any('؀' <= ch <= 'ۿ' for ch in s):
                c.drawRightString(x_right, y, s)
            else:
                c.drawString(x_left, y, s)
            y -= line_h
            if y < 20 * mm:
                c.showPage()
                c.setFont(chosen_font, 12)
                y = height - 20 * mm
    c.save()


def cmd_export(args: argparse.Namespace) -> None:
    os.makedirs(args.outdir, exist_ok=True)
    msgs = load_messages(args.inp)
    md = to_markdown(msgs, title=args.title)
    formats = {f.strip().lower() for f in args.formats.split(',') if f.strip()}
    if 'md' in formats:
        save_markdown(md, os.path.join(args.outdir, 'transcript.md'))
        logging.info('تم إنشاء Markdown')
    if 'docx' in formats:
        try:
            save_docx(msgs, os.path.join(args.outdir, 'transcript.docx'), title=args.title, font_name=args.font)
            logging.info('تم إنشاء DOCX')
        except Exception as e:
            logging.exception('فشل DOCX: %s', e)
    if 'xlsx' in formats:
        try:
            save_xlsx(msgs, os.path.join(args.outdir, 'transcript.xlsx'))
            logging.info('تم إنشاء XLSX')
        except Exception as e:
            logging.exception('فشل XLSX: %s', e)
    if 'pdf' in formats:
        try:
            extra_fonts = [args.pdf_font] if args.pdf_font else None
            save_pdf(md, os.path.join(args.outdir, 'transcript.pdf'), font_paths=extra_fonts)
            logging.info('تم إنشاء PDF')
        except Exception as e:
            logging.exception('فشل PDF: %s', e)
    print('تم التصدير إلى', args.outdir)


# ============================================================
# القسم (ب): تدريب وتقييم Causal LM
# ============================================================

def build_datasets(
    train_files: List[str],
    valid_files: Optional[List[str]],
    valid_ratio: float,
    tokenizer,
    block_size: int,
    test_files: Optional[List[str]] = None,
    test_ratio: Optional[float] = None,
):
    from datasets import DatasetDict, load_dataset

    data_files = {"train": train_files}
    if valid_files:
        data_files["validation"] = valid_files
    if test_files:
        data_files["test"] = test_files

    if valid_files or test_files:
        raw = load_dataset("text", data_files=data_files)
        dsd = DatasetDict(raw)
    else:
        raw_all = load_dataset("text", data_files={"train": train_files})
        split = raw_all["train"].train_test_split(test_size=valid_ratio, seed=42)
        train_part, valid_part = split["train"], split["test"]
        if test_ratio and test_ratio > 0:
            split2 = train_part.train_test_split(test_size=test_ratio, seed=42)
            train_part, test_part = split2["train"], split2["test"]
            dsd = DatasetDict({"train": train_part, "validation": valid_part, "test": test_part})
        else:
            dsd = DatasetDict({"train": train_part, "validation": valid_part})

    def tokenize_function(examples):
        return tokenizer(examples["text"])  # بلا padding ثابت هنا

    tokenized = dsd.map(tokenize_function, batched=True, remove_columns=["text"])  # type: ignore[arg-type]

    def group_texts(examples):
        concatenated = {k: sum(examples[k], []) for k in examples.keys()}
        total_length = len(concatenated["input_ids"])
        total_length = (total_length // block_size) * block_size
        result = {
            k: [t[i:i + block_size] for i in range(0, total_length, block_size)]
            for k, t in concatenated.items()
        }
        result["labels"] = result["input_ids"].copy()
        return result

    tokenized = tokenized.map(group_texts, batched=True, num_proc=1)
    return tokenized


class SampleGenerationCallback:
    """Callback بسيط لتوليد عينات نصية في نهاية كل حقبة."""

    def __init__(
        self,
        tokenizer,
        prompts: List[str],
        outdir: str,
        max_length: int = 128,
        temperature: float = 0.8,
        top_p: float = 0.9,
        top_k: int = 50,
        num_return_sequences: int = 1,
        no_repeat_ngram_size: int = 2,
    ):
        from transformers import TrainerCallback

        class _CB(TrainerCallback):
            def __init__(self, outer: "SampleGenerationCallback"):
                self.outer = outer

            def on_epoch_end(self, args, state, control, **kwargs):  # type: ignore[override]
                import torch

                model = kwargs.get('model')
                tokenizer = self.outer.tokenizer
                outdir = self.outer.outdir
                os.makedirs(outdir, exist_ok=True)
                if model is None:
                    return
                device = model.device
                gen_kwargs = dict(
                    max_length=self.outer.max_length,
                    do_sample=True,
                    temperature=self.outer.temperature,
                    top_p=self.outer.top_p,
                    top_k=self.outer.top_k,
                    no_repeat_ngram_size=self.outer.no_repeat_ngram_size,
                    num_return_sequences=self.outer.num_return_sequences,
                    pad_token_id=tokenizer.eos_token_id,
                )
                for pi, prompt in enumerate(self.outer.prompts, 1):
                    inputs = tokenizer(prompt, return_tensors='pt').to(device)
                    with torch.no_grad():
                        out = model.generate(**inputs, **gen_kwargs)
                    texts = [tokenizer.decode(seq, skip_special_tokens=True) for seq in out]
                    epoch_id = int(state.epoch) if state.epoch is not None else 0
                    path = os.path.join(outdir, f"epoch_{epoch_id:02d}_prompt_{pi}.txt")
                    with open(path, 'w', encoding='utf-8') as f:
                        f.write(f"# Prompt:\n{prompt}\n\n# Outputs:\n\n")
                        for i, t in enumerate(texts, 1):
                            f.write(f"=== {i} ===\n{t}\n\n")

        self.tokenizer = tokenizer
        self.prompts = prompts
        self.outdir = outdir
        self.max_length = max_length
        self.temperature = temperature
        self.top_p = top_p
        self.top_k = top_k
        self.num_return_sequences = num_return_sequences
        self.no_repeat_ngram_size = no_repeat_ngram_size
        self.cb = _CB(self)


def cmd_train(args: argparse.Namespace) -> None:
    from transformers import (
        AutoModelForCausalLM,
        AutoTokenizer,
        DataCollatorForLanguageModeling,
        EarlyStoppingCallback,
        Trainer,
        TrainingArguments,
    )

    os.makedirs(args.output_dir, exist_ok=True)

    tokenizer = AutoTokenizer.from_pretrained(args.model_name, use_fast=True)
    if tokenizer.pad_token is None:
        tokenizer.pad_token = tokenizer.eos_token

    tokenized = build_datasets(
        train_files=args.train_files,
        valid_files=args.valid_files,
        valid_ratio=args.valid_ratio,
        tokenizer=tokenizer,
        block_size=args.block_size,
        test_files=args.test_files,
        test_ratio=args.test_ratio,
    )

    model = AutoModelForCausalLM.from_pretrained(args.model_name)
    if getattr(model.config, "pad_token_id", None) is None:
        model.config.pad_token_id = tokenizer.pad_token_id

    data_collator = DataCollatorForLanguageModeling(tokenizer=tokenizer, mlm=False)

    eval_strategy = "steps" if args.eval_steps and args.eval_steps > 0 else "epoch"
    save_strategy = "steps" if args.save_steps and args.save_steps > 0 else eval_strategy

    training_args = TrainingArguments(
        output_dir=args.output_dir,
        overwrite_output_dir=True,
        num_train_epochs=args.epochs,
        per_device_train_batch_size=args.batch_size,
        gradient_accumulation_steps=args.grad_accum,
        learning_rate=args.lr,
        weight_decay=args.weight_decay,
        lr_scheduler_type=args.lr_scheduler_type,
        warmup_ratio=args.warmup_ratio,
        fp16=args.fp16,
        bf16=args.bf16,
        save_steps=args.save_steps if save_strategy == "steps" else None,
        eval_steps=args.eval_steps if eval_strategy == "steps" else None,
        save_total_limit=args.save_total_limit,
        logging_steps=50,
        report_to=["tensorboard"],
        logging_dir=os.path.join(args.output_dir, "tb_logs"),
        evaluation_strategy=eval_strategy,
        save_strategy=save_strategy,
        load_best_model_at_end=True,
        metric_for_best_model="eval_loss",
        greater_is_better=False,
        seed=args.seed,
        gradient_checkpointing=args.gradient_checkpointing,
        eval_accumulation_steps=args.eval_accumulation_steps,
        push_to_hub=args.push_to_hub,
        hub_model_id=args.hub_model_id,
    )

    def compute_metrics(eval_pred):
        preds, labels = eval_pred
        if hasattr(preds, 'ndim') and preds.ndim == 3:
            import numpy as np

            preds = preds.argmax(-1)
        mask = labels != -100
        correct = (preds[mask] == labels[mask]).sum()
        total = mask.sum()
        accuracy = float(correct) / float(total) if total > 0 else 0.0
        return {"accuracy": accuracy}

    trainer = Trainer(
        model=model,
        args=training_args,
        train_dataset=tokenized["train"],
        eval_dataset=tokenized.get("validation"),
        tokenizer=tokenizer,
        data_collator=data_collator,
        compute_metrics=compute_metrics,
    )

    if args.early_stopping_patience and args.early_stopping_patience > 0:
        trainer.add_callback(EarlyStoppingCallback(early_stopping_patience=args.early_stopping_patience))

    sample_prompts: List[str] = []
    if args.sample_prompts_file and os.path.isfile(args.sample_prompts_file):
        with open(args.sample_prompts_file, encoding='utf-8') as f:
            for ln in f:
                ln = ln.strip()
                if ln:
                    sample_prompts.append(ln)
    if args.sample_prompt:
        sample_prompts.extend(args.sample_prompt)

    if sample_prompts:
        sgc = SampleGenerationCallback(
            tokenizer=tokenizer,
            prompts=sample_prompts,
            outdir=os.path.join(args.output_dir, 'samples'),
            max_length=args.sample_max_length,
            temperature=args.sample_temperature,
            top_p=args.sample_top_p,
            top_k=args.sample_top_k,
            num_return_sequences=args.sample_num_return_sequences,
            no_repeat_ngram_size=args.sample_no_repeat_ngram_size,
        )
        trainer.add_callback(sgc.cb)

    trainer.train()

    eval_metrics = trainer.evaluate()
    eval_loss = eval_metrics.get("eval_loss")
    ppl = math.exp(eval_loss) if isinstance(eval_loss, (int, float)) else float('nan')
    eval_metrics["perplexity"] = ppl

    trainer.save_model(args.output_dir)
    tokenizer.save_pretrained(args.output_dir)

    metrics_val_path = os.path.join(args.output_dir, "metrics_val.json")
    with open(metrics_val_path, "w", encoding="utf-8") as f:
        json.dump({k: (float(v) if isinstance(v, (int, float)) else v) for k, v in eval_metrics.items()}, f, ensure_ascii=False, indent=2)

    if tokenized.get("test") is not None:
        test_metrics = trainer.evaluate(eval_dataset=tokenized["test"])
        test_loss = test_metrics.get("eval_loss")
        test_ppl = math.exp(test_loss) if isinstance(test_loss, (int, float)) else float('nan')
        test_metrics["perplexity"] = test_ppl
        metrics_test_path = os.path.join(args.output_dir, "metrics_test.json")
        with open(metrics_test_path, "w", encoding="utf-8") as f:
            json.dump({k: (float(v) if isinstance(v, (int, float)) else v) for k, v in test_metrics.items()}, f, ensure_ascii=False, indent=2)
        logging.info("تم حفظ مقاييس الاختبار إلى %s", metrics_test_path)
    else:
        test_metrics = None

    print("\n📊 Validation Metrics:")
    for k, v in eval_metrics.items():
        print(f"- {k}: {v}")
    if test_metrics is not None:
        print("\n🧪 Test Metrics:")
        for k, v in test_metrics.items():
            print(f"- {k}: {v}")
    print("\n✅ انتهى التدريب. أفضل نموذج محفوظ في:", args.output_dir)
    print("📁 مقاييس التحقق:", metrics_val_path)
    if test_metrics is not None:
        print("📁 مقاييس الاختبار:", metrics_test_path)
    print("🧾 عينات التوليد (إن وُجدت):", os.path.join(args.output_dir, 'samples'))
    print("🧪 TensorBoard logs:", os.path.join(args.output_dir, "tb_logs"))


# ============================================================
# القسم (ج): توليد سيناريوهات
# ============================================================

def cmd_generate(args: argparse.Namespace) -> None:
    import torch
    from transformers import AutoModelForCausalLM, AutoTokenizer

    model_path = args.model_path or args.model_name
    tokenizer = AutoTokenizer.from_pretrained(model_path)
    model = AutoModelForCausalLM.from_pretrained(model_path)

    if tokenizer.pad_token is None:
        tokenizer.pad_token = tokenizer.eos_token
    model.config.pad_token_id = tokenizer.pad_token_id

    device = "cuda" if torch.cuda.is_available() else "cpu"
    model.to(device)

    inputs = tokenizer(args.prompt, return_tensors="pt").to(device)

    gen_kwargs = dict(
        max_length=args.max_length,
        do_sample=not args.use_beam_search,
        temperature=args.temperature,
        top_p=args.top_p,
        top_k=args.top_k,
        repetition_penalty=args.repetition_penalty,
        no_repeat_ngram_size=args.no_repeat_ngram_size,
        num_return_sequences=args.num_return_sequences,
        pad_token_id=tokenizer.eos_token_id,
    )
    if args.use_beam_search:
        gen_kwargs.update(dict(num_beams=args.num_beams, early_stopping=True))

    with torch.no_grad():
        output = model.generate(**inputs, **gen_kwargs)

    texts = [tokenizer.decode(seq, skip_special_tokens=True) for seq in output]
    print("\n📝 المخرجات:")
    for i, t in enumerate(texts, 1):
        print(f"\n=== # {i} ===\n{t}")


# ============================================================
# القسم (د): مثال دوال spells (اختياري)
# ============================================================

def is_powerful(magic: str) -> bool:
    powerful_set = {"Sourcery", "More Sourcery"}
    return magic in powerful_set


def find_more(magicks: List[str]) -> List[str]:
    return [m for m in magicks if is_powerful(m)]


def print_all(spells: List[str]) -> None:
    for s in spells:
        print(s)


def cmd_spells(args: argparse.Namespace) -> None:
    items = [s.strip() for s in (args.list or '').split(',') if s.strip()]
    strong = find_more(items)
    print("✨ Powerful spells:")
    print_all(strong)


# ============================================================
# القسم (هـ): واجهة Gradio (مستشار العلاقات والصحة الجنسية)
# ============================================================

def _ascii_ratio(s: str) -> float:
    if not s:
        return 0.0
    letters = sum(ch in ascii_letters for ch in s)
    digits_count = sum(ch in digits for ch in s)
    return (letters + digits_count) / max(len(s), 1)


AR_KEYWORDS = [
    "جنس", "حميمي", "حميمية", "علاقة", "علاقات", "زوج", "زوجة", "زواج", "طلاق",
    "انفصال", "خيانة", "غيرة", "تعارف", "مواعدة", "تواصل", "حوار", "اتصال",
    "موافقة", "رضا", "حدود", "ثقة", "قرب", "مشاعر", "حميم", "رغبة", "متعة",
    "نشوة", "اورجازم", "نقاش", "توافق", "قابلية", "خلل جنسي", "ضعف جنسي",
    "قذف", "انتصاب", "تشحيم", "مزلق", "صحة جنسية", "تثقيف جنسي", "استمناء",
    "تخيلات", "خيال جنسي", "آمنة", "ممارسة آمنة", "علاج زوجي", "علاج جنسي",
    "ارشاد", "استشارة", "علاج نفسي"
]
EN_KEYWORDS = [
    "sex", "sexual", "intimacy", "relationship", "marriage", "divorce", "dating",
    "consent", "boundaries", "trust", "communication", "orgasm", "pleasure",
    "arousal", "erectile", "ejaculation", "libido", "vibrator", "toy", "lube",
    "sex education", "sex therapy", "couples therapy", "anxiety"
]
ALL_PHRASES = sorted({*AR_KEYWORDS, *EN_KEYWORDS}, key=len, reverse=True)
KEYWORD_PATTERN = re.compile(r"\b(" + "|".join(map(re.escape, ALL_PHRASES)) + r")\b", re.IGNORECASE)
AR_LETTERS = re.compile(r"[؀-ۿ]")


def is_relevant(question: str) -> bool:
    q = (question or "").strip()
    if not q:
        return False
    has_ar = bool(AR_LETTERS.search(q))
    has_en_like = _ascii_ratio(q) >= 0.25
    relevant = bool(KEYWORD_PATTERN.search(q))
    return relevant and (has_ar or has_en_like)


def cmd_serve(args: argparse.Namespace) -> None:
    try:
        import gradio as gr
    except Exception as e:
        raise RuntimeError("Gradio غير مثبت. ثبّت gradio أولاً: pip install gradio") from e
    try:
        from huggingface_hub import InferenceClient
    except Exception as e:
        raise RuntimeError("huggingface_hub غير مثبت. ثبّت الحزمة: pip install huggingface_hub") from e

    api_token = os.getenv(args.api_token_env or "")
    if not api_token:
        raise RuntimeError(f"مفقود التوكن في المتغيّر البيئي '{args.api_token_env}'.")

    client = InferenceClient(args.model_id, token=api_token)

    def _chat_with_retries(messages, max_tokens, temperature, top_p, retries=3, delay=0.8):
        last_error = None
        for attempt in range(1, retries + 1):
            try:
                stream = client.chat_completion(
                    messages=messages,
                    max_tokens=max_tokens,
                    stream=True,
                    temperature=temperature,
                    top_p=top_p,
                )
                for chunk in stream:
                    token = None
                    try:
                        token = chunk.choices[0].delta.content  # type: ignore[attr-defined]
                    except Exception:
                        pass
                    if token:
                        yield token
                return
            except Exception as e:
                last_error = e
                if attempt < retries:
                    time.sleep(delay * attempt)
                else:
                    raise last_error

    def respond(message, history, system_message, max_tokens, temperature, top_p):
        history = history or []
        if len(history) == 0:
            yield {"role": "assistant", "content": "مرحبًا. اسأل سؤالك بالعربية عن العلاقات أو الصحة الجنسية. سأجيب باختصار ووضوح."}
            return
        if not is_relevant(message):
            yield {"role": "assistant", "content": "من فضلك اسأل عن العلاقات، التواصل، الموافقة، أو الصحة الجنسية. سأجيب بالعربية."}
            return
        messages = [{"role": "system", "content": system_message}]
        messages.extend(history)
        messages.append({"role": "user", "content": message})
        response_text = ""
        try:
            for token in _chat_with_retries(messages, max_tokens, temperature, top_p):
                response_text += token
                yield {"role": "assistant", "content": response_text}
        except Exception as e:  # pragma: no cover - network errors
            yield {"role": "assistant", "content": f"حدث خطأ في الاتصال بالنموذج: {type(e).__name__}. حاول لاحقًا."}

    demo = gr.ChatInterface(
        fn=respond,
        additional_inputs=[
            gr.Textbox(
                value="أنت مستشار علاقات وصحة جنسية محترف. أجب بالعربية بوضوح وحزم وبلا أحكام. قدّم توازنًا بين العلم والعملية.",
                label="رسالة النظام",
            ),
            gr.Slider(minimum=50, maximum=1024, value=512, step=1, label="أقصى عدد رموز", visible=False),
            gr.Slider(minimum=0.1, maximum=1.0, value=0.7, step=0.05, label="درجة العشوائية", visible=False),
            gr.Slider(minimum=0.1, maximum=1.0, value=0.9, step=0.05, label="Top-p", visible=False),
        ],
        chatbot=gr.Chatbot(type="messages", label="مستشار علاقات وصحة جنسية"),
        title="مستشار العلاقات والصحة الجنسية",
        description="اكتب سؤالك بالعربية. الإجابات معلوماتية ومباشرة.",
    )
    demo.launch(server_name=args.host, server_port=args.port, show_api=False)


# CLI رئيسي
# ============================================================

def build_parser() -> argparse.ArgumentParser:
    ap = argparse.ArgumentParser(description="أداة موحّدة: تصدير دردشات + تدريب/تقييم + توليد")
    sub = ap.add_subparsers(dest="command", required=True)

    # export
    ape = sub.add_parser("export", help="تصدير سجلات الدردشة")
    ape.add_argument("--in", dest="inp", required=True, help="ملف الإدخال .csv أو .jsonl أو .json")
    ape.add_argument("--out", dest="outdir", required=True, help="مجلد الإخراج")
    ape.add_argument("--title", dest="title", default="سجل الدردشة", help="عنوان الوثائق")
    ape.add_argument("--formats", default="md,docx,xlsx,pdf", help="md,docx,xlsx,pdf")
    ape.add_argument("--pdf-font", dest="pdf_font", default=None, help="مسار خط TTF لملف PDF")
    ape.add_argument("--font", dest="font", default="Arial", help="اسم خط DOCX")
    ape.set_defaults(func=cmd_export)

    # train
    apt = sub.add_parser("train", help="تدريب Causal LM مع تقييم")
    apt.add_argument("--train_files", nargs="+", required=True)
    apt.add_argument("--valid_files", nargs="+", default=None)
    apt.add_argument("--valid_ratio", type=float, default=0.05)
    apt.add_argument("--test_files", nargs="+", default=None)
    apt.add_argument("--test_ratio", type=float, default=0.0)
    apt.add_argument("--model_name", default="EleutherAI/gpt-neo-125M")
    apt.add_argument("--output_dir", default="./results")
    apt.add_argument("--block_size", type=int, default=256)
    apt.add_argument("--epochs", type=float, default=3.0)
    apt.add_argument("--batch_size", type=int, default=8)
    apt.add_argument("--grad_accum", type=int, default=1)
    apt.add_argument("--lr", type=float, default=5e-5)
    apt.add_argument("--weight_decay", type=float, default=0.01)
    apt.add_argument("--lr_scheduler_type", default="linear")
    apt.add_argument("--warmup_ratio", type=float, default=0.0)
    apt.add_argument("--fp16", action="store_true")
    apt.add_argument("--bf16", action="store_true")
    apt.add_argument("--save_steps", type=int, default=0)
    apt.add_argument("--eval_steps", type=int, default=0)
    apt.add_argument("--save_total_limit", type=int, default=2)
    apt.add_argument("--seed", type=int, default=42)
    apt.add_argument("--early_stopping_patience", type=int, default=0)
    apt.add_argument("--gradient_checkpointing", action="store_true")
    apt.add_argument("--eval_accumulation_steps", type=int, default=1)
    apt.add_argument("--push_to_hub", action="store_true")
    apt.add_argument("--hub_model_id", default=None)
    apt.add_argument("--sample_prompts_file", default=None, help="ملف نصي يحوي prompt في كل سطر")
    apt.add_argument("--sample_prompt", action='append', default=None, help="يمكن تمرير هذا الخيار عدة مرات لإضافة برومبتات")
    apt.add_argument("--sample_max_length", type=int, default=128)
    apt.add_argument("--sample_temperature", type=float, default=0.8)
    apt.add_argument("--sample_top_p", type=float, default=0.9)
    apt.add_argument("--sample_top_k", type=int, default=50)
    apt.add_argument("--sample_num_return_sequences", type=int, default=1)
    apt.add_argument("--sample_no_repeat_ngram_size", type=int, default=2)

    apt.set_defaults(func=cmd_train)

    # generate
    apg = sub.add_parser("generate", help="توليد نص من نموذج (محلي أو Hub)")
    apg.add_argument("--prompt", required=True)
    apg.add_argument("--model_path", default=None, help="مسار مجلد نموذج محفوظ محليًا")
    apg.add_argument("--model_name", default="EleutherAI/gpt-neo-125M", help="اسم نموذج من Hub إذا لم يوجد model_path")
    apg.add_argument("--max_length", type=int, default=128)
    apg.add_argument("--temperature", type=float, default=0.8)
    apg.add_argument("--top_p", type=float, default=0.9)
    apg.add_argument("--top_k", type=int, default=50)
    apg.add_argument("--repetition_penalty", type=float, default=1.1)
    apg.add_argument("--no_repeat_ngram_size", type=int, default=2)
    apg.add_argument("--num_return_sequences", type=int, default=1)
    apg.add_argument("--use_beam_search", action="store_true")
    apg.add_argument("--num_beams", type=int, default=4)
    apg.set_defaults(func=cmd_generate)

    # spells (اختياري)
    aps = sub.add_parser("spells", help="مثال بسيط لدوال مساعدة")
    aps.add_argument("--list", default="Sourcery,Alchemy,More Sourcery")
    aps.set_defaults(func=cmd_spells)

    # serve (Gradio)
    apv = sub.add_parser("serve", help="تشغيل واجهة Gradio للمستشار")
    apv.add_argument("--model_id", default="HuggingFaceH4/zephyr-7b-beta")
    apv.add_argument("--api_token_env", default="seksuoloog_sexologist_bot", help="اسم متغير البيئة الذي يحتوي التوكن")
    apv.add_argument("--host", default="0.0.0.0")
    apv.add_argument("--port", type=int, default=7860)
    apv.set_defaults(func=cmd_serve)

    return ap


def main():
    parser = build_parser()
    args = parser.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
